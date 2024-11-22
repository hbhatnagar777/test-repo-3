from __future__ import unicode_literals
import datetime
import random
import string
import os
import time
import shutil
import wikipedia
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas


class Wiki:
    """Class to interact with wikipedia API to fetch random articles"""

    def __init__(self, log):
        """Initializes the wiki object

                Args:

                    log (object) -- instance of the logger module

                Returns:

                    object  --  Instance of Wiki class

        """
        self.log = log
        self.wiki = wikipedia

    def _fetch_random_article(self):
        _title = self.wiki.random()
        try:
            _page = self.wiki.WikipediaPage(_title)
            return _page
        except wikipedia.exceptions.DisambiguationError:
            # if random article caused some disambiguation, fetch another instead
            return self._fetch_random_article()

    def fetch_message_from_wiki(self, unicode_data=False):
        """Fetches page title and content of a random page from wikipedia

                Args:

                    unicode_data (boolean)  --  True if unicode data has to be fetched

                                                Default: False

                Returns:

                    Dict of following type: {'subject': <some subject>, 'body': <page content>}

        """
        message = {}
        try:

            self.__get_wikiapi_service(unicode_data)
            _article = self._fetch_random_article()

            message['subject'] = _article.original_title
            self.log.info(u'Message subject: %s', message['subject'])

            message['body'] = _article.content

            return message
        except Exception as excp:
            self.log.exception('Exception while creating message from wiki')
            raise Exception(str(excp))

    def fetch_subject_from_wiki(self, unicode_data: bool = False):

        self.__get_wikiapi_service(unicode_data)
        _article = self._fetch_random_article()

        _subject = _article.original_title
        self.log.info(u'Subject: %s', _subject)
        return _subject

    def fetch_text_from_wiki(self, unicode_data: bool = False):
        self.__get_wikiapi_service(unicode_data)
        _article = self._fetch_random_article()
        return _article.content

    def fetch_summary_from_wiki(self, unicode: bool = False):
        self.__get_wikiapi_service(unicode_data=unicode)
        _article = self._fetch_random_article()

        _summary = _article.summary
        return _summary

    def __get_wikiapi_service(self, unicode_data=False):
        """Returns a random language wiki api service. Useful to create unicode mails.

                Args:

                    unicode_data (boolean):

                        Default: False, if true it returns a wiki service other than English

                Returns:

                    wiki_api_service (str)  --  Wikipedia API endpoint

        """

        if unicode_data:
            prefix = random.choices(list(wikipedia.languages().keys()))
            self.log.info("Using language with prefix: {}".format(prefix))
            self.wiki.set_lang(prefix)
        else:
            self.wiki.set_lang('eng')


class O365DataGenerator:
    """
        Method to generate test data for O365/ D365 Agents

    """

    def __init__(self, logger_obj):
        self.log = logger_obj
        self.__datetime = datetime.datetime.utcnow()
        self.wiki = Wiki(log=self.log)

    def gen_alias_name(self, special_chars: bool = False):
        _alias = ''.join(random.choices(string.ascii_lowercase, k=7))
        if special_chars:
            _special_char = ''.join(random.choices("!# $%&'*+-/=?^_`{|}~", k=2))
            #   special chars cannot be first or last letter of the name, so generate and append to alias

            _alias = _alias + _special_char

        alias_name = (f'{_alias}'
                      f'{self.testcase_id}')
        return alias_name

    @staticmethod
    def gen_display_name(special_chars: bool = False):
        if special_chars:
            _display_name = ''.join(random.choices(string.printable, k=5))
        else:
            _display_name = ''.join(random.choices(string.ascii_lowercase, k=7))
        return _display_name

    def generate_email_message(self, unicode: bool = False):
        _msg = self.wiki.fetch_message_from_wiki(unicode_data=unicode)
        return _msg

    def gen_file(self,
                 no_of_docs=5,
                 unicode_data=False,
                 pdf=False,
                 doc_path: str = "C:\\tmp"):
        """This method creates a new docx document and saves it into the directory
            defined in constants file.

                Args:

                    doc_path:
                    no_of_docs (int)  --  Number of docx documents to create

                    unicode_data (boolean)  --  True if unicode document has to be created

                                                Default: False

                    pdf  (boolean)  --  True if pdf files need to be created

        """
        try:

            if os.path.exists(doc_path):
                time.sleep(5)
                shutil.rmtree(doc_path)

            os.makedirs(doc_path)

            for _ in range(no_of_docs):
                from docx import Document
                document = Document()
                wiki_message = self.wiki.fetch_message_from_wiki(unicode_data)
                file_name = self._make_valid_filename(wiki_message['subject'])
                document.add_heading(file_name)
                paragraph = document.add_paragraph(wiki_message['body'])
                doc_format = paragraph.paragraph_format

                filename = os.path.join(
                    doc_path,
                    f'{file_name}.docx')
                document.save(filename)

                if pdf:
                    wiki_message = self.wiki.fetch_message_from_wiki(unicode_data)
                    file_name = self._make_valid_filename(wiki_message['subject'])
                    out_file = os.path.join(
                        doc_path,
                        f'{file_name}.pdf')

                    can = canvas.Canvas(out_file, pagesize=letter)
                    width, height = letter
                    text_object = can.beginText()
                    text_object.setTextOrigin(width - 19 * cm, height - 2.5 * cm)
                    text_object.setFont("Helvetica-Oblique", 14)
                    text_object.textLines(wiki_message['body'])
                    can.drawText(text_object)
                    can.showPage()
                    can.save()

        except Exception as excp:
            self.log.exception('Exception while creating documents.')
            raise Exception(str(excp))

    def _make_valid_filename(self, file_name):
        """This method returns the valid file name for Windows OS
        after removing invalid characters.

                Args:

                    file_name  (str)  --  Name of the file

                Returns:

                    file_name  (str)  --  Returns a valid filename

        """

        # file_name = f'file{str(int(time.time()))}'
        file_name = ''.join(c for c in file_name if c in string.ascii_letters)
        self.log.info('Valid filename for the file: %s', file_name)
        return file_name

    def delete_file(self, file_name):
        """Deletes the specified file

            Args:

                file_name (str)      --      name of the file

        """
        self.log.info("Deleting File with name: {}".format(file_name))
        if os.path.exists(file_name):
            os.remove(file_name)
        else:
            raise Exception("The file does not exist")

    def modify_file(self, file_name):
        """Updates the specified file

            Args:

                file_name (str)      --      name of the file

        """
        self.log.info("Modifying file with name: {}".format(file_name))
        file = open(file_name, "a")
        size = random.randint(50, 100)
        chars = ''.join([random.choice(string.ascii_letters) for i in range(size)]) + "\n"
        file.write(chars)
        file.close()

    @staticmethod
    def generate_test_value_for_list_item_column(field_type_kind):
        """Generates test value for the list item based on column type

            Args:

                field_type_kind (int)       --      field id of column i.e., type of the column

        """
        size = random.randint(50, 100)
        if field_type_kind == 1:
            return size
        elif field_type_kind == 2:
            return ''.join([random.choice(string.ascii_letters) for i in range(size)])
        elif field_type_kind == 3:
            return ''.join([random.choice(string.ascii_letters) for i in range(size)]) + "\n" + \
                   ''.join([random.choice(string.ascii_letters) for i in range(size)])

    def gen_email_addr(self):
        """
        Generates a random email address

        Returns:
            list: data
        """

        return f"{self.__gen_string(8, string.ascii_lowercase)}@testemail.com"

    def gen_url(self):
        """
        Generate a URL

        Returns:
            list: data
        """
        return f"https://www.{self.__gen_string(5, string.ascii_lowercase)}.com"

    def gen_date(self):
        """
        Generates a random date

        Returns:
            list: data
        """
        return self.__datetime.strftime("%Y-%m-%d")

    def gen_datetime(self):
        """
        Generates a date time string
        """
        return self.__datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')

    def gen_name(self, use_unicode: bool = False):
        """
            Generate a name
        """
        if use_unicode:
            return "{} {} {}".format(self.get_random_unicode(10), self.get_random_unicode(10),
                                     self.get_random_unicode(5))
        return "{} {} {}".format(self.__gen_string(10), self.__gen_string(10), self.__gen_string(5))

    @staticmethod
    def gen_boolean():
        """
        Generates random boolean data
        """
        return random.random() > 0.5

    @staticmethod
    def gen_decimal():
        """
            Generate a random decimal
        """
        return random.random()

    def gen_text(self, use_unicode: bool = False):
        """
        Generates short (length 10) string

        Args:
            use_unicode     (bool): Generate unicode text
        Returns:
            data
        """
        if use_unicode:
            return self.get_random_unicode(10)
        return self.__gen_string(10)

    def gen_long_text(self, unicode: bool = False, length: int = 1000):
        """
        Generates long (length 100) string

        Args:
            unicode    (bool): Whether to have unicode data
            length      (int):  Length of text to be generated
        Returns:
            str: data
        """
        if unicode:
            return self.get_random_unicode(length=length)
        return self.wiki.fetch_subject_from_wiki(unicode_data=unicode)

    def gen_number(self):
        """
        Generates number
        """
        return self.__gen_string(random.randint(1, 10), string.digits)

    @staticmethod
    def gen_double():
        """
        Generates double values
        """
        return round(random.uniform(10, 10 ** 4), 2)

    @staticmethod
    def gen_percent():
        """
        Generates percentage figures

        """
        return str(random.randint(0, 100))

    def gen_website_url(self):
        """
            Generate a website URL for test purpose
        """
        return "https://{}.{}".format(self.__gen_string(10), self.__gen_string(3))

    def gen_phone(self):
        """
        Generates list of random phone numbers

        Args:
            rec_count (int): Number of records

        Returns:
            list: data
        """
        return self.__gen_string(10, string.digits)

    def __gen_string(self, str_length: int, str_type=string.ascii_letters):
        """
        Generates a random string

        Args:
            str_length (int): Length of string
            str_type (StringType): Use this parameter to control the type of characters in the return string

        Returns:
            str: Random string
        """
        return ''.join(random.choices(str_type, k=str_length))

    @staticmethod
    def get_random_unicode(length=25):
        """
            Generates a random string of unicode characters
        """
        get_char = chr

        include_ranges = [
            (0x0021, 0x0021),
            (0x0023, 0x0026),
            (0x0028, 0x007E),
            (0x00A1, 0x00AC),
            (0x00AE, 0x00FF),
            (0x0100, 0x017F),
            (0x0180, 0x024F),
            (0x2C60, 0x2C7F),
            (0x16A0, 0x16F0),
            (0x0370, 0x0377),
            (0x037A, 0x037E),
            (0x0384, 0x038A),
            (0x038C, 0x038C),
        ]

        alphabet = [
            get_char(code_point) for current_range in include_ranges
            for code_point in range(current_range[0], current_range[1] + 1)
        ]
        return ''.join(random.choice(alphabet) for i in range(length))
